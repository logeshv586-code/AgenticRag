"""
Document Parser — Handles extraction from PDF, TXT, DOCX, CSV, HTML, Markdown,
Images (LLM vision + OCR fallback), Audio (Whisper transcription), and Multilingual docs.
"""
import os
import logging

logger = logging.getLogger(__name__)

OLLAMA_BASE_URL = "http://localhost:11434"

# ═══════════════════════════════════════════════════════════
#  Optional Imports — graceful degradation
# ═══════════════════════════════════════════════════════════

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image, ImageFilter, ImageEnhance
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from PIL import Image as PILImage
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from langdetect import detect as detect_language
    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False

try:
    from deep_translator import GoogleTranslator
    TRANSLATOR_AVAILABLE = True
except ImportError:
    TRANSLATOR_AVAILABLE = False

# Check Whisper availability (faster-whisper preferred, fall back to openai-whisper)
WHISPER_AVAILABLE = False
WHISPER_BACKEND = None
try:
    from faster_whisper import WhisperModel
    WHISPER_AVAILABLE = True
    WHISPER_BACKEND = "faster-whisper"
except ImportError:
    try:
        import whisper
        WHISPER_AVAILABLE = True
        WHISPER_BACKEND = "openai-whisper"
    except ImportError:
        pass


# ═══════════════════════════════════════════════════════════
#  Main Entry Point
# ═══════════════════════════════════════════════════════════

def parse_document(file_path: str) -> str:
    """
    Reads an uploaded file and extracts text content.
    Supports: PDF, TXT, DOCX, CSV, HTML, MD, Images (LLM vision/OCR), Audio (Whisper).
    All non-English text is auto-detected and translated to English for indexing.
    """
    if not os.path.exists(file_path):
        return "File not found."

    ext = os.path.splitext(file_path)[1].lower()
    extracted_text = ""

    try:
        if ext == '.pdf':
            extracted_text = _parse_pdf(file_path)
        elif ext == '.txt':
            extracted_text = _parse_text(file_path)
        elif ext == '.docx':
            extracted_text = _parse_docx(file_path)
        elif ext == '.csv':
            extracted_text = _parse_csv(file_path)
        elif ext in ('.html', '.htm'):
            extracted_text = _parse_html(file_path)
        elif ext in ('.md', '.markdown'):
            extracted_text = _parse_markdown(file_path)
        elif ext in ('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp'):
            extracted_text = _parse_image(file_path)
        elif ext in ('.mp3', '.wav', '.m4a', '.ogg', '.webm', '.flac'):
            extracted_text = _transcribe_audio(file_path)
        else:
            return f"Unsupported file type: {ext}"

        if not extracted_text:
            return f"No text extracted from {ext} file."

        # Auto-translate non-English content
        extracted_text = _ensure_english(extracted_text.strip(), file_path)
        return extracted_text

    except Exception as e:
        logger.error(f"Error parsing document ({ext}): {e}")
        return f"Error parsing document ({ext}): {str(e)}"


# ═══════════════════════════════════════════════════════════
#  Document Parsers
# ═══════════════════════════════════════════════════════════

def _parse_pdf(file_path: str) -> str:
    """Extract text from PDF, falling back to OCR for image-based PDFs."""
    if not PDF_AVAILABLE:
        return "PDF library (pdfplumber) is not installed."

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
            elif OCR_AVAILABLE:
                img = page.to_image(resolution=300).original
                text_parts.append(pytesseract.image_to_string(img))

    # Also extract tables
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    clean_row = [str(cell or '') for cell in row]
                    text_parts.append(' | '.join(clean_row))

    return '\n'.join(text_parts)


def _parse_text(file_path: str) -> str:
    """Read plain text files with encoding detection."""
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    return "Could not decode text file."


def _parse_docx(file_path: str) -> str:
    """Extract text from DOCX files."""
    if not DOCX_AVAILABLE:
        return "DOCX library (python-docx) is not installed."

    doc = docx.Document(file_path)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append(' | '.join(cells))

    return '\n'.join(text_parts)


def _parse_csv(file_path: str) -> str:
    """Extract text from CSV files with intelligent column handling."""
    import csv

    text_parts = []
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        reader = csv.reader(f)
        headers = None
        for i, row in enumerate(reader):
            if i == 0:
                headers = row
                text_parts.append('Columns: ' + ', '.join(row))
            else:
                if headers:
                    pairs = [f"{h}: {v}" for h, v in zip(headers, row) if v.strip()]
                    text_parts.append('; '.join(pairs))
                else:
                    text_parts.append(', '.join(row))
            if i > 500:
                text_parts.append(f"... (truncated at {i} rows)")
                break

    return '\n'.join(text_parts)


def _parse_html(file_path: str) -> str:
    """Extract text from local HTML files."""
    from bs4 import BeautifulSoup

    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')

    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()

    content = soup.find('article') or soup.find('main') or soup.body or soup
    return content.get_text(separator='\n', strip=True)


def _parse_markdown(file_path: str) -> str:
    """Extract text from Markdown files (strip formatting)."""
    import re

    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()

    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    text = re.sub(r'\[(.+?)\]\(.*?\)', r'\1', text)

    return text


# ═══════════════════════════════════════════════════════════
#  Image Understanding (LLM Vision + OCR Fallback)
# ═══════════════════════════════════════════════════════════

def _parse_image(file_path: str) -> str:
    """
    Understand images using:
    1. Ollama vision model (llava/gemma3) for semantic understanding — rich description
    2. Fallback: pytesseract OCR for text extraction
    """
    # Try LLM Vision first (much richer than OCR)
    vision_result = _describe_image_with_llm(file_path)
    if vision_result:
        return vision_result

    # Fallback to OCR
    return _parse_image_ocr(file_path)


def _describe_image_with_llm(file_path: str) -> str:
    """Send image to Ollama vision model (llava/gemma3) for rich description."""
    try:
        import requests
        import base64

        # Check if a vision model is available
        resp = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=3)
        if resp.status_code != 200:
            return ""

        models = [m["name"] for m in resp.json().get("models", [])]
        vision_models = [m for m in models if any(v in m.lower() for v in ["llava", "gemma3", "bakllava", "moondream", "llama3.2-vision"])]

        if not vision_models:
            logger.info("No Ollama vision model found, falling back to OCR")
            return ""

        vision_model = vision_models[0]

        # Encode image to base64
        with open(file_path, "rb") as f:
            image_b64 = base64.b64encode(f.read()).decode("utf-8")

        # Send to Ollama vision API
        resp = requests.post(
            f"{OLLAMA_BASE_URL}/api/generate",
            json={
                "model": vision_model,
                "prompt": "Describe this image in detail. Extract all text visible in the image. Describe any charts, diagrams, tables, or visual elements. Be thorough and comprehensive for search indexing.",
                "images": [image_b64],
                "stream": False
            },
            timeout=120
        )
        if resp.status_code == 200:
            description = resp.json().get("response", "")
            if description:
                logger.info(f"Image described by {vision_model}: {len(description)} chars")
                return f"[Image Analysis by {vision_model}]\n{description}"

    except Exception as e:
        logger.warning(f"LLM vision failed: {e}")

    return ""


def _parse_image_ocr(file_path: str) -> str:
    """Extract text from images using OCR with preprocessing (fallback)."""
    if not OCR_AVAILABLE:
        if not PIL_AVAILABLE:
            return "Image parsing requires Pillow. Install with: pip install Pillow"
        return "OCR requires pytesseract. Install Tesseract OCR from https://github.com/tesseract-ocr/tesseract"

    image = Image.open(file_path)
    image = image.convert('L')
    image = ImageEnhance.Contrast(image).enhance(2.0)
    image = image.filter(ImageFilter.SHARPEN)

    if os.name == 'nt':
        tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path

    text = pytesseract.image_to_string(image, config='--oem 3 --psm 6')
    return f"[OCR Text]\n{text}"


# ═══════════════════════════════════════════════════════════
#  Audio Transcription (Local Whisper — No Cloud Needed)
# ═══════════════════════════════════════════════════════════

def _transcribe_audio(file_path: str) -> str:
    """
    Transcribe audio files using local Whisper model.
    Supports: MP3, WAV, M4A, OGG, WebM, FLAC.
    Uses faster-whisper (preferred) or openai-whisper as fallback.
    """
    if not WHISPER_AVAILABLE:
        return (
            "Audio transcription requires faster-whisper. "
            "Install with: pip install faster-whisper\n"
            "Or: pip install openai-whisper"
        )

    try:
        # Convert non-WAV audio to WAV using pydub if needed
        ext = os.path.splitext(file_path)[1].lower()
        audio_path = file_path

        if ext not in ('.wav',):
            audio_path = _convert_audio_to_wav(file_path)

        if WHISPER_BACKEND == "faster-whisper":
            return _transcribe_faster_whisper(audio_path)
        else:
            return _transcribe_openai_whisper(audio_path)

    except Exception as e:
        logger.error(f"Audio transcription failed: {e}")
        return f"Audio transcription failed: {str(e)}"
    finally:
        # Clean up converted temp file
        if 'audio_path' in locals() and audio_path != file_path and os.path.exists(audio_path):
            os.remove(audio_path)


def _convert_audio_to_wav(file_path: str) -> str:
    """Convert audio file to WAV format using pydub."""
    try:
        from pydub import AudioSegment
        import tempfile

        ext = os.path.splitext(file_path)[1].lower().strip('.')
        if ext == 'm4a':
            ext = 'mp4'
        if ext == 'ogg':
            ext = 'ogg'

        audio = AudioSegment.from_file(file_path, format=ext)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
        audio.export(tmp.name, format='wav')
        tmp.close()
        return tmp.name
    except Exception as e:
        logger.warning(f"Audio conversion failed, trying direct: {e}")
        return file_path  # Try original file directly


def _transcribe_faster_whisper(audio_path: str) -> str:
    """Transcribe using faster-whisper (ctranslate2-based, very fast)."""
    from faster_whisper import WhisperModel

    logger.info(f"Transcribing {audio_path} with faster-whisper...")
    # Use 'base' model by default (fast + accurate). 'small' or 'medium' for better accuracy.
    model = WhisperModel("base", device="cpu", compute_type="int8")
    segments, info = model.transcribe(audio_path, beam_size=5)

    text_parts = [seg.text.strip() for seg in segments]
    full_text = " ".join(text_parts)

    lang = info.language if hasattr(info, 'language') else "unknown"
    logger.info(f"Transcribed {len(full_text)} chars, detected language: {lang}")
    return f"[Audio Transcript — Language: {lang}]\n{full_text}"


def _transcribe_openai_whisper(audio_path: str) -> str:
    """Transcribe using openai-whisper (fallback)."""
    import whisper

    logger.info(f"Transcribing {audio_path} with openai-whisper...")
    model = whisper.load_model("base")
    result = model.transcribe(audio_path)
    text = result.get("text", "")
    lang = result.get("language", "unknown")
    logger.info(f"Transcribed {len(text)} chars, detected language: {lang}")
    return f"[Audio Transcript — Language: {lang}]\n{text}"


# ═══════════════════════════════════════════════════════════
#  Multilingual Support — Auto Detect & Translate
# ═══════════════════════════════════════════════════════════

def _ensure_english(text: str, file_path: str = "") -> str:
    """
    Auto-detect language and translate non-English content to English.
    Keeps both original + translated text for richer indexing.
    """
    if not text or len(text.strip()) < 20:
        return text

    if not LANGDETECT_AVAILABLE:
        return text  # No detection available, return as-is

    try:
        detected_lang = detect_language(text[:500])  # Sample first 500 chars

        if detected_lang == 'en':
            return text  # Already English

        logger.info(f"Detected non-English content ({detected_lang}) in {os.path.basename(file_path)}, translating...")

        if not TRANSLATOR_AVAILABLE:
            return f"[Original Language: {detected_lang}]\n{text}"

        # Translate in chunks (Google Translate has char limits)
        translated = _translate_chunked(text, detected_lang)
        return (
            f"[Original Language: {detected_lang}]\n"
            f"[Original Text]\n{text[:2000]}{'...' if len(text) > 2000 else ''}\n\n"
            f"[English Translation]\n{translated}"
        )

    except Exception as e:
        logger.warning(f"Language detection/translation failed: {e}")
        return text


def _translate_chunked(text: str, source_lang: str, chunk_size: int = 4500) -> str:
    """Translate text in chunks to avoid API limits."""
    if len(text) <= chunk_size:
        return GoogleTranslator(source=source_lang, target='en').translate(text)

    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    translated_chunks = []
    for chunk in chunks[:10]:  # Cap at 10 chunks (~45k chars)
        try:
            t = GoogleTranslator(source=source_lang, target='en').translate(chunk)
            translated_chunks.append(t)
        except Exception as e:
            logger.warning(f"Chunk translation failed: {e}")
            translated_chunks.append(chunk)
    return " ".join(translated_chunks)


# ═══════════════════════════════════════════════════════════
#  Supported Extensions
# ═══════════════════════════════════════════════════════════

def get_supported_extensions() -> list:
    """Return list of supported file extensions."""
    exts = ['.pdf', '.txt', '.csv', '.html', '.htm', '.md', '.markdown']
    if DOCX_AVAILABLE:
        exts.append('.docx')
    # Images — always try (LLM vision or OCR)
    exts.extend(['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp'])
    # Audio — always list (will gracefully fail if whisper not installed)
    exts.extend(['.mp3', '.wav', '.m4a', '.ogg', '.webm', '.flac'])
    return exts
